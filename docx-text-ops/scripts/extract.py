"""Extract all text from a Word (.docx) file for inspection and validation.

Prints body text, table cells, headers, footers, footnotes, endnotes,
comments, and text boxes to stdout, one paragraph per line,
so the output can be piped to grep, awk, or other tools.

Usage:
    python extract.py <file.docx>

Examples:
    python extract.py document.docx
    python extract.py edited.docx | grep -iE "xxxx|lorem|ipsum"
    python extract.py document.docx | grep "\\[footnote"
"""

import argparse
import sys
from pathlib import Path

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

DOCX_EXTENSIONS = {".docx", ".docm", ".dotx", ".dotm"}

_FOOTNOTES_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
_ENDNOTES_RT  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
_COMMENTS_RT  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"


def _runs_text(element) -> str:
    """Concatenate all <w:t> text inside element, respecting xml:space=preserve."""
    parts = []
    for t in element.iter(qn("w:t")):
        parts.append(t.text or "")
    return "".join(parts).strip()


def extract(input_file: str) -> str:
    src = Path(input_file)
    if not src.exists():
        return f"Error: {input_file} does not exist"
    if src.suffix.lower() not in DOCX_EXTENSIONS:
        return f"Error: {input_file} must be a .docx / .docm / .dotx / .dotm file"

    try:
        doc = docx.Document(str(src))
        lines = []

        # Tracked-changes warning
        body_el = doc.element.body
        has_ins = next(body_el.iter(qn("w:ins")), None) is not None
        has_del = next(body_el.iter(qn("w:del")), None) is not None
        if has_ins or has_del:
            lines.append(
                "<!-- WARNING: document contains tracked changes. "
                "Accept or reject all changes in Word before editing. -->"
            )

        # Body paragraphs
        lines.append("<!-- body -->")
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables
        for i, table in enumerate(doc.tables, 1):
            lines.append(f"<!-- table {i} -->")
            seen = set()
            for row in table.rows:
                for cell in row.cells:
                    # python-docx may yield merged cells multiple times
                    if id(cell) in seen:
                        continue
                    seen.add(id(cell))
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            lines.append(f"[table {i}] {text}")

        # Headers and footers
        hf_pairs = [
            ("header", lambda s: s.header),
            ("footer", lambda s: s.footer),
            ("first_page_header", lambda s: s.first_page_header),
            ("first_page_footer", lambda s: s.first_page_footer),
            ("even_page_header", lambda s: s.even_page_header),
            ("even_page_footer", lambda s: s.even_page_footer),
        ]
        hf_seen = set()
        for i, section in enumerate(doc.sections, 1):
            for hf_name, getter in hf_pairs:
                try:
                    hf = getter(section)
                    if hf is None or hf.is_linked_to_previous:
                        continue
                    hf_id = id(hf._element)
                    if hf_id in hf_seen:
                        continue
                    hf_seen.add(hf_id)
                    for para in hf.paragraphs:
                        text = para.text.strip()
                        if text:
                            lines.append(f"[{hf_name}] {text}")
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    text = para.text.strip()
                                    if text:
                                        lines.append(f"[{hf_name} table] {text}")
                except Exception:
                    continue

        # Text boxes in body (<w:txbxContent>)
        txbx_found = False
        for txbx in body_el.iter(qn("w:txbxContent")):
            if not txbx_found:
                lines.append("<!-- text boxes -->")
                txbx_found = True
            for p in txbx.findall(f".//{qn('w:p')}"):
                text = _runs_text(p)
                if text:
                    lines.append(f"[textbox] {text}")

        # Footnotes
        try:
            fn_part = doc.part.part_related_by(_FOOTNOTES_RT)
            fn_root = fn_part._element
            fn_items = [
                fn for fn in fn_root.findall(f".//{qn('w:footnote')}")
                if fn.get(qn("w:id")) not in ("-1", "0")
            ]
            if fn_items:
                lines.append("<!-- footnotes -->")
                for fn in fn_items:
                    fn_id = fn.get(qn("w:id"), "?")
                    for p in fn.findall(f".//{qn('w:p')}"):
                        text = _runs_text(p)
                        if text:
                            lines.append(f"[footnote {fn_id}] {text}")
        except KeyError:
            pass

        # Endnotes
        try:
            en_part = doc.part.part_related_by(_ENDNOTES_RT)
            en_root = en_part._element
            en_items = [
                en for en in en_root.findall(f".//{qn('w:endnote')}")
                if en.get(qn("w:id")) not in ("-1", "0")
            ]
            if en_items:
                lines.append("<!-- endnotes -->")
                for en in en_items:
                    en_id = en.get(qn("w:id"), "?")
                    for p in en.findall(f".//{qn('w:p')}"):
                        text = _runs_text(p)
                        if text:
                            lines.append(f"[endnote {en_id}] {text}")
        except KeyError:
            pass

        # Comments
        try:
            cm_part = doc.part.part_related_by(_COMMENTS_RT)
            cm_root = cm_part._element
            cm_items = cm_root.findall(f".//{qn('w:comment')}")
            if cm_items:
                lines.append("<!-- comments -->")
                for cm in cm_items:
                    cm_id = cm.get(qn("w:id"), "?")
                    author = cm.get(qn("w:author"), "")
                    for p in cm.findall(f".//{qn('w:p')}"):
                        text = _runs_text(p)
                        if text:
                            lines.append(f"[comment {cm_id} by {author}] {text}")
        except KeyError:
            pass

        return "\n".join(lines)
    except Exception as e:
        return f"Error: {e}"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from a Word document")
    parser.add_argument("input_file", help="Word file to read (.docx or .docm)")
    args = parser.parse_args()

    result = extract(args.input_file)
    print(result)
    if result.startswith("Error"):
        sys.exit(1)
